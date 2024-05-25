import os
import boto3
import fitz
from docx import Document
from dotenv import load_dotenv
import streamlit as st
from openai import OpenAI

# Load environment variables from .env file
load_dotenv()

# Initialize Textract client
aws_access_key_id = os.getenv('AWS_ACCESS_KEY_ID')
aws_secret_access_key = os.getenv('AWS_SECRET_ACCESS_KEY')
aws_region = os.getenv('AWS_REGION', 'ap-south-1')  # Default to ap-south-1 if not provided

client = boto3.client('textract',
                      aws_access_key_id=aws_access_key_id,
                      aws_secret_access_key=aws_secret_access_key,
                      region_name=aws_region)


ai_client = OpenAI(
    # This is the default and can be omitted
    api_key=os.environ.get("OPENAI_API_KEY"),
)


def extract_text_from_pdf(pdf_data):
    st.write("Extracting text from PDF...")
    try:
        pdf_document = fitz.open(stream=pdf_data, filetype='pdf')
        text = ""
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            text += page.get_text("text")
        st.success("Text extracted successfully!")
        return text
    except Exception as e:
        st.error(f"Error extracting text: {e}")
        return None

def create_docx_from_text(text, output_file):
    st.write(f"Creating DOCX file: {output_file}...")
    try:
        doc = Document()
        doc.add_heading('Extracted Text', level=1)

        # Add paragraphs with preserved formatting
        for para in text.split('\n\n'):
            doc.add_paragraph(para)

        doc.save(output_file)
        st.success(f"DOCX file generated: {output_file}")
        return output_file
    except Exception as e:
        st.error(f"Error creating DOCX file: {e}")
        return None

def generate_resume_from_text(text):
    st.write("Generating resume from extracted text using OpenAI...")

    # OpenAI GPT-3.5 model for text completion
    model = "gpt-4o"

    try:
        response = ai_client.chat.completions.create(
            messages=[
                {"role": "user", "content": text}
            ],
            model=model
        )
        generated_resume = response.choices[0].message.content.strip()


        st.success("Resume generated successfully!")
        st.write("extracted data",generated_resume)
        return generated_resume
    except Exception as e:
        st.error(f"Error generating resume: {e}")
        return None

def save_text_to_pdf(text, output_file):
    st.write(f"Saving resume as PDF: {output_file}...")
    try:
        doc = Document()
        doc.add_heading('Generated Resume', level=1)

        # Add paragraphs with preserved formatting
        for para in text.split('\n\n'):
            doc.add_paragraph(para)

        doc.save(output_file)
        st.success(f"PDF file saved: {output_file}")
    except Exception as e:
        st.error(f"Error saving PDF file: {e}")

def main(file):
    if file is not None:
        # Read the uploaded file
        pdf_data = file.read()
        file.seek(0)  # Reset file pointer

        # Extract text from PDF using PyMuPDF
        extracted_text = extract_text_from_pdf(pdf_data)

        if extracted_text:
            # Generate resume from extracted text using OpenAI
            generated_resume = generate_resume_from_text(extracted_text)

            if generated_resume:
                # Save generated resume as PDF
                output_file = "generated_resume.pdf"
                save_text_to_pdf(generated_resume, output_file)

                # Provide download link for the generated PDF
                st.markdown("### Download Generated Resume PDF")
                st.markdown(get_binary_file_downloader_html(output_file, file_label="Generated Resume PDF"), unsafe_allow_html=True)

def get_binary_file_downloader_html(bin_file, file_label='File'):
    with open(bin_file, 'rb') as f:
        base64_pdf = base64.b64encode(f.read()).decode('utf-8')
    return f'<a href="data:application/octet-stream;base64,{base64_pdf}" download="{file_label}.pdf">Click here to download {file_label}</a>'



if __name__ == "__main__":
    st.title("PDF to Resume PDF Converter")

    # File uploader
    uploaded_file = st.file_uploader("Upload a PDF file", type=['pdf'])

    if st.button("Convert to Resume PDF") and uploaded_file:
        main(uploaded_file)
